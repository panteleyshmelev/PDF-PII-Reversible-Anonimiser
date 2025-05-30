# app/routes/documents.py

import uuid
import json
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, PlainTextResponse
from app.core import masking
import os
import fitz  # PyMuPDF for text extraction from PDF
import docx  # For .docx files
from collections import defaultdict

router = APIRouter()

# Define new directory paths
UPLOAD_DIR = "data/uploads"
MASKED_TXTS_DIR = "data/masked_txts"
UNMASKED_TXTS_DIR = "data/unmasked_txts"
MAPPING_FILE = "data/masking_maps.json"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(MASKED_TXTS_DIR, exist_ok=True)
os.makedirs(UNMASKED_TXTS_DIR, exist_ok=True)

def load_mappings():
    if os.path.exists(MAPPING_FILE):
        try:
            with open(MAPPING_FILE, "r", encoding="utf-8") as f:
                content = f.read()
                if not content:
                    return {}
                return json.loads(content)
        except json.JSONDecodeError:
            return {}
    return {}

def save_mappings(mappings):
    with open(MAPPING_FILE, "w", encoding="utf-8") as f:
        json.dump(mappings, f, indent=4)

@router.post("/process-to-masked-txt", tags=["Text Processing"])
async def process_to_masked_txt(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    original_file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")
    masked_txt_path = os.path.join(MASKED_TXTS_DIR, f"masked_{file_id}.txt")

    with open(original_file_path, "wb") as buffer:
        buffer.write(await file.read())

    full_text_parts = [] # Use a list to gather text parts
    file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''

    try:
        if file_extension == "pdf":
            doc_pdf = fitz.open(original_file_path)
            for page in doc_pdf:
                full_text_parts.append(page.get_text())
            doc_pdf.close()
        elif file_extension == "txt":
            with open(original_file_path, "r", encoding="utf-8") as f_txt:
                full_text_parts.append(f_txt.read())
        elif file_extension == "docx":
            try:
                doc_word = docx.Document(original_file_path)
                # Extract text from paragraphs
                for para in doc_word.paragraphs:
                    full_text_parts.append(para.text)
                
                # --- START: ADDED TABLE TEXT EXTRACTION ---
                for table in doc_word.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text_parts.append(cell.text) # Add text from each cell
                # --- END: ADDED TABLE TEXT EXTRACTION ---

            except Exception as e_docx:
                raise HTTPException(status_code=400, detail=f"Error processing .docx file. Ensure it's a valid .docx format. (Details: {str(e_docx)})")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_extension}. Please upload PDF, TXT, or DOCX.")

        full_text = "\n".join(full_text_parts) # Join all parts with a newline

        if not full_text.strip():
            raise HTTPException(status_code=400, detail="No text content found or extracted from the uploaded file.")

    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        if os.path.exists(original_file_path):
            os.remove(original_file_path)
        raise HTTPException(status_code=500, detail=f"Error extracting text from file '{file.filename}': {str(e)}")

    # --- PII Detection and Filtering (remains the same) ---
    pii_entities_raw = masking.get_pii_entities(full_text)
    pii_entities_raw.sort(key=lambda e: (e.start, -e.end))

    filtered_entities = []
    for current_entity in pii_entities_raw:
        is_subset = False
        for kept_entity in filtered_entities:
            if current_entity.start >= kept_entity.start and \
               current_entity.end <= kept_entity.end:
                is_subset = True
                break
        if not is_subset:
            filtered_entities.append(current_entity)

    # --- Placeholder Generation and Mapping (remains the same) ---
    text_replacements_map = {}
    unmasking_map = {}
    counters = defaultdict(int)
    
    unique_texts_and_types = {}
    for entity in filtered_entities:
        original_text_segment = full_text[entity.start:entity.end]
        if original_text_segment not in unique_texts_and_types:
            unique_texts_and_types[original_text_segment] = entity.entity_type

    for original_text, entity_type in unique_texts_and_types.items():
        count = counters[entity_type]
        placeholder = f"<{entity_type}_{count}>"
        text_replacements_map[original_text] = placeholder
        unmasking_map[placeholder] = original_text
        counters[entity_type] += 1
    
    all_mappings = load_mappings()
    all_mappings[file_id] = unmasking_map
    save_mappings(all_mappings)

    # --- Masking the Extracted Text (String Replacement - remains the same) ---
    masked_text = full_text
    sorted_original_texts = sorted(text_replacements_map.keys(), key=len, reverse=True)
    for orig_text in sorted_original_texts:
        placeholder = text_replacements_map[orig_text]
        masked_text = masked_text.replace(orig_text, placeholder)

    # --- Save Masked Text to .txt file (remains the same) ---
    with open(masked_txt_path, "w", encoding="utf-8") as f_out:
        f_out.write(masked_text)

    if os.path.exists(original_file_path):
        os.remove(original_file_path)

    return {
        "file_id": file_id,
        "original_filename": file.filename,
        "masked_txt_file_path": masked_txt_path,
        "detail": f"File '{file.filename}' processed, text extracted (including table content for DOCX), PII masked, and saved as .txt."
    }

# The /demask-txt endpoint remains unchanged
@router.get("/demask-txt/{file_id}", tags=["Text Processing"])
async def demask_txt(file_id: str):
    # ... (rest of the function is the same)
    masked_txt_path = os.path.join(MASKED_TXTS_DIR, f"masked_{file_id}.txt")
    unmasked_txt_path = os.path.join(UNMASKED_TXTS_DIR, f"unmasked_{file_id}.txt")

    if not os.path.exists(masked_txt_path):
        raise HTTPException(status_code=404, detail=f"Masked text file not found for ID: {file_id}")

    all_mappings = load_mappings()
    if file_id not in all_mappings:
        raise HTTPException(status_code=404, detail=f"Unmasking map not found for ID: {file_id}")

    unmasking_map_for_file = all_mappings[file_id]

    with open(masked_txt_path, "r", encoding="utf-8") as f_in:
        masked_text = f_in.read()

    unmasked_text = masked_text
    sorted_placeholders = sorted(unmasking_map_for_file.keys(), key=len, reverse=True)
    for placeholder in sorted_placeholders:
        original_text = unmasking_map_for_file[placeholder]
        unmasked_text = unmasked_text.replace(placeholder, original_text)

    with open(unmasked_txt_path, "w", encoding="utf-8") as f_out:
        f_out.write(unmasked_text)
        
    return PlainTextResponse(content=unmasked_text)